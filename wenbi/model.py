import dspy
import os
import re
import sys
import json
import subprocess
import traceback
from datetime import datetime
from docx import Document
from redlines import Redlines
import google.generativeai as genai
from wenbi.utils import segment


def get_lm_config(model_string, base_url=None):
    """
    Determine provider and return config dict for dspy.LM.
    Supports: ollama, openai, gemini (google-genai)
    """
    if not model_string:
        # Default to Ollama
        return {
            "base_url": "http://localhost:11434",
            "model": "ollama/qwen3",
        }
    parts = model_string.strip().split("/")
    provider = parts[0].lower() if parts else ""
    if provider == "ollama":
        return {
            "base_url": base_url or "http://localhost:11434",
            "model": model_string,
        }
    elif provider == "openai":
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise ValueError("OPENAI_API_KEY environment variable not set.")
        return {
            "model": model_string.replace("openai/", ""),
            "api_key": api_key,
        }
    elif provider == "gemini":
        # For Gemini models, use our custom adapter
        try:
            # Make sure the required package is installed
            import google.generativeai

            # Extract model name from the model string
            model_name = model_string.replace("gemini/", "")
            print(f"Creating adapter for Gemini model: {model_name}")

            # Create a custom DSPy adapter for Gemini (preferred)
            # This adapter implements the DSPy LM interface directly
            try:
                lm = CustomDSPyGeminiAdapter(
                    model_name=model_name,
                    temperature=0.1,
                    max_tokens=2048
                )
                return {"lm": lm}
            except Exception as adapter_error:
                print(f"Error creating CustomDSPyGeminiAdapter, falling back to legacy adapter: {adapter_error}")
                # Fall back to legacy adapter if the DSPy adapter fails
                adapter = GeminiAdapter(model_name)
                return {
                    "model": model_name,
                    "adapter": adapter,
                }
        except ImportError:
            raise ImportError("Please install Google Generative AI: pip install google-generativeai")
        except Exception as e:
            print(f"Error setting up Gemini model: {e}")
            raise ValueError(f"Failed to initialize Gemini model: {str(e)}")
    else:
        # Unknown provider, fallback to user input
        return {"base_url": base_url, "model": model_string}


def is_ollama(model_string):
    """
    Check the model_string (if provided) for the provider.
    If the model string starts with "ollama/", return "http://localhost:11434".
    If model_string is empty, default to "ollama/qwen3" with base_url "http://localhost:11434".
    Otherwise, return None.
    """
    if not model_string:
        return "http://localhost:11434"  # default for empty input
    parts = model_string.strip().split("/")
    if parts and parts[0].lower() == "ollama":
        return "http://localhost:11434"
    return None


def translate(
    vtt_path,
    output_dir=None,
    translate_language="Chinese",
    llm="",
    chunk_length=8,
    max_tokens=50000,
    timeout=3600,
    temperature=0.1,
    base_url="http://localhost:11434",
):
    """
    Translate English VTT content to a bilingual markdown file using the target language provided.

    Args:
        vtt_path (str): Path to the English VTT file
        output_dir (str): Directory for output files
        translate_language (str): Target language for translation
        llm (str): LLM model identifier
        chunk_length (int): Number of sentences per chunk for segmentation
        max_tokens (int): Maximum number of tokens for the LLM
        timeout (int): Timeout for the LLM in seconds
        temperature (float): Temperature for the LLM
        base_url (str): Base URL for the LLM

    Returns:
        str: Path to the generated markdown file
    """
    segmented_text = segment(vtt_path, sentence_count=chunk_length)
    paragraphs = segmented_text.split("\n\n")

    model_id = llm if llm else "ollama/qwen3"
    lm_config = get_lm_config(model_id, base_url=base_url)
    lm_config["max_tokens"] = max_tokens
    lm_config["timeout_s"] = timeout
    lm_config["temperature"] = temperature
    lm = dspy.LM(**lm_config)
    dspy.configure(lm=lm)

    class Translate(dspy.Signature):
        english_text = dspy.InputField(desc="English text to translate")
        translated_text = dspy.OutputField(
            desc=f"Translation into {translate_language}"
        )

    translator = dspy.ChainOfThought(Translate)
    translated_pairs = []

    for para in paragraphs:
        if para.strip():
            response = translator(english_text=para)
            translated_pairs.append(
                f"# English\n{para}\n\n# {translate_language}\n{response.translated_text}\n\n---\n"
            )

    markdown_content = "\n".join(translated_pairs)
    output_file = os.path.splitext(vtt_path)[0] + "_bilingual.md"
    with open(output_file, "w", encoding="utf-8") as f:
        f.write(markdown_content)

    return output_file


def rewrite(
    file_path,
    output_dir=None,
    llm="",
    rewrite_lang="Chinese",
    chunk_length=8,
    max_tokens=50000,
    timeout=3600,
    temperature=0.1,
    base_url="http://localhost:11434",
):
    """
    Rewrites text by first segmenting the file into paragraphs.

    Args:
        file_path (str): Path to the input file
        output_dir (str, optional): Output directory
        llm (str): LLM model identifier
        rewrite_lang (str): Target language for rewriting (default: Chinese)
        chunk_length (int): Number of sentences per chunk for segmentation
        max_tokens (int): Maximum number of tokens for the LLM
        timeout (int): Timeout for the LLM in seconds
        temperature (float): Temperature for the LLM
        base_url (str): Base URL for the LLM
    """
    segmented_text = segment(file_path, sentence_count=chunk_length)
    paragraphs = segmented_text.split("\n\n")

    model_id = llm if llm else "ollama/qwen3"
    lm_config = get_lm_config(model_id, base_url=base_url)
    lm_config["max_tokens"] = max_tokens
    lm_config["timeout_s"] = timeout
    lm_config["temperature"] = temperature
    lm = dspy.LM(**lm_config)
    dspy.configure(lm=lm)

    rewritten_paragraphs = []
    for para in paragraphs:
        class ParaRewrite(dspy.Signature):
            """
            Rewrite this text in {rewrite_lang} from oral to written.  Follow these rules strictly:
            1. Correct any basic grammar, punctuation, or usage errors.
            2. Improve clarity while preserving the original meaning and scholarly tone (trying your best not to change the structure of sentence)
            3. IMPORTANT: Maintaining the original meaning and length (97% of original)
            """
            text: str = dspy.InputField(
                desc=f"Spoken text to rewrite in {rewrite_lang}"
            )
            rewritten: str = dspy.OutputField(
                desc=f"Rewritten paragraph in {rewrite_lang}"
            )
        rewrite = dspy.ChainOfThought(ParaRewrite)
        response = rewrite(text=para)
        rewritten_paragraphs.append(response.rewritten)

    rewritten_text = "\n\n".join(rewritten_paragraphs)
    if output_dir:
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        out_file = os.path.join(output_dir, f"{base_name}_rewritten.md")
        with open(out_file, "w", encoding="utf-8") as f:
            f.write(rewritten_text)
    else:
        out_file = None

    return rewritten_text


# Custom adapter for Google Gemini
# Custom DSPy adapter for Gemini models
class CustomDSPyGeminiAdapter(dspy.LM):
    def __init__(self, model_name, **kwargs):
        super().__init__(model=model_name)  # Pass model name to parent class
        # Using globally imported genai module
        global genai

        self.model_name = model_name
        self.kwargs = kwargs

        # First try with API key (preferred method)
        self.api_key = os.getenv("GOOGLE_API_KEY")

        if self.api_key:
            # Configure with explicit API key
            print(f"Using API key authentication for Gemini model: {model_name}")
            genai.configure(api_key=self.api_key)
        else:
            # Set path to application default credentials
            print(f"No API key found, setting ADC credentials path")
            os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "/home/ajiap/.config/gcloud/application_default_credentials.json"

            # Don't configure genai with project_id as it causes issues
            # It will use ADC automatically when no API key is provided

        try:
            self.model = genai.GenerativeModel(model_name)
            print(f"Successfully initialized Gemini model: {model_name}")
        except Exception as e:
            print(f"Error initializing Gemini model: {e}")
            raise

    def basic_request(self, prompt, **kwargs):
        """Send a basic completion request to the model."""
        try:
            # Process and send the prompt to Gemini
            generation_config = {
                "temperature": kwargs.get("temperature", self.kwargs.get("temperature", 0.1)),
                "top_p": kwargs.get("top_p", 0.95),
                "top_k": kwargs.get("top_k", 40),
                "max_output_tokens": kwargs.get("max_tokens", self.kwargs.get("max_tokens", 1024)),
            }

            safety_settings = [
                {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
                {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
                {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
                {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
            ]

            print(f"Sending prompt to Gemini ({len(prompt)} chars)")
            response = self.model.generate_content(
                prompt,
                generation_config=generation_config,
                safety_settings=safety_settings
            )

            # Check if response has text property
            if hasattr(response, 'text'):
                response_text = response.text
                print(f"Received valid response from Gemini ({len(response_text)} chars)")
                return response_text
            else:
                print(f"Warning: Response has no text attribute: {response}")
                return "No response text was generated. Please try again."

        except Exception as e:
            error_msg = str(e)
            print(f"Gemini API error: {error_msg}")

            # Check for specific error types
            if "PERMISSION_DENIED" in error_msg or "Permission denied" in error_msg:
                return "Error: Permission denied accessing Google Gemini API. Please check your API key or credentials."
            elif "project" in error_msg and ("not found" in error_msg or "invalid" in error_msg):
                return "Error: Invalid Google Cloud project. Please check your GOOGLE_CLOUD_PROJECT environment variable."
            elif "quota" in error_msg:
                return "Error: Quota exceeded for Google Gemini API. Please try again later."
            else:
                return f"Error from Gemini API: {error_msg}"

    def _complete(self, prompt, **kwargs):
        """Complete a prompt with the Gemini model."""
        response_text = self.basic_request(prompt, **kwargs)

        # Format response to match DSPy expectations
        return {
            "choices": [{
                "message": {
                    "content": response_text
                }
            }]
        }

    def generate(self, prompt, **kwargs):
        """Generate text from a prompt."""
        response = self._complete(prompt, **kwargs)
        # Process result for DSPy
        try:
            if isinstance(response, dict) and "choices" in response:
                return response["choices"][0]["message"]["content"]
            elif isinstance(response, str):
                return response
            elif isinstance(response, list) and len(response) > 0:
                if isinstance(response[0], dict) and "message" in response[0]:
                    return response[0]["message"].get("content", "No content found")
                elif isinstance(response[0], str):
                    return response[0]
            return "Error processing model response: Unexpected format"
        except (KeyError, IndexError, TypeError) as e:
            print(f"Error processing response: {e}, response type: {type(response)}")
            if isinstance(response, dict):
                return str(response)
            elif isinstance(response, list) and len(response) > 0:
                return str(response[0])
            return "Error processing model response"

    def __call__(self, prompt, **kwargs):
        """Make the adapter callable for compatibility with DSPy."""
        return self._complete(prompt, **kwargs)

# Legacy adapter for backward compatibility
class GeminiAdapter:
    def __init__(self, model_name):
        # Using globally imported genai module
        global genai

        self.model_name = model_name

        # First try with API key (preferred method)
        self.api_key = os.getenv("GOOGLE_API_KEY")

        if self.api_key:
            # Configure with explicit API key
            print(f"Using API key authentication for Gemini model: {model_name}")
            genai.configure(api_key=self.api_key)
        else:
            # Set path to application default credentials
            print(f"No API key found, setting ADC credentials path")
            os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "/home/ajiap/.config/gcloud/application_default_credentials.json"

            # Don't configure genai with project_id as it causes issues
            # It will use ADC automatically when no API key is provided

        try:
            self.model = genai.GenerativeModel(model_name)
            print(f"Successfully initialized Gemini model: {model_name}")
        except Exception as e:
            print(f"Error initializing Gemini model: {e}")
            raise

    def __call__(self, prompt, **kwargs):
        try:
            # Process and send the prompt to Gemini
            generation_config = {
                "temperature": kwargs.get("temperature", 0.1),
                "top_p": 0.95,
                "top_k": 40,
                "max_output_tokens": kwargs.get("max_tokens", 1024),
            }

            safety_settings = [
                {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
                {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
                {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
                {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
            ]

            print(f"Sending prompt to Gemini ({len(prompt)} chars)")
            response = self.model.generate_content(
                prompt,
                generation_config=generation_config,
                safety_settings=safety_settings
            )

            # Check if response has text property
            if hasattr(response, 'text'):
                print(f"Received valid response from Gemini ({len(response.text)} chars)")
                return {
                    "choices": [{
                        "message": {
                            "content": response.text
                        }
                    }]
                }
            else:
                print(f"Warning: Response has no text attribute: {response}")
                # Handle empty or invalid response
                return {
                    "choices": [{
                        "message": {
                            "content": "No response text was generated. Please try again."
                        }
                    }]
                }

        except Exception as e:
            error_msg = str(e)
            print(f"Gemini API error: {error_msg}")

            # Check for specific error types
            if "PERMISSION_DENIED" in error_msg or "Permission denied" in error_msg:
                error_response = "Error: Permission denied accessing Google Gemini API. Please check your API key or credentials."
            elif "project" in error_msg and ("not found" in error_msg or "invalid" in error_msg):
                error_response = "Error: Invalid Google Cloud project. Please check your GOOGLE_CLOUD_PROJECT environment variable."
            elif "quota" in error_msg:
                error_response = "Error: Quota exceeded for Google Gemini API. Please try again later."
            else:
                error_response = f"Error from Gemini API: {error_msg}"

            # Return a structured error message
            return {
                "choices": [{
                    "message": {
                        "content": error_response
                    }
                }]
            }




def academic(
    file_path,
    output_dir=None,
    llm="",
    academic_llm="",  # Added academic_llm parameter
    academic_lang="English",
    chunk_length=8,
    max_tokens=250000,
    timeout=3600,
    temperature=0.1,
    base_url="http://localhost:11434",
):
    """Rewrites text in academic style while preserving markdown formatting and footnotes.
    Returns the academic text without writing to a file unless output_dir is specified."""
    print(f"DEBUG: academic function called with file_path={file_path}, llm={llm}, academic_llm={academic_llm}")

    # Type checking and defensive coding
    if not isinstance(file_path, str):
        print(f"WARNING: file_path is not a string: {type(file_path)}")
        file_path = str(file_path)

    if not isinstance(llm, str):
        print(f"WARNING: llm is not a string: {type(llm)}")
        llm = str(llm) if llm is not None else ""

    if not isinstance(academic_llm, str):
        print(f"WARNING: academic_llm is not a string: {type(academic_llm)}")
        academic_llm = str(academic_llm) if academic_llm is not None else ""

    # Read markdown content
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Split content into blocks while preserving markdown and footnote elements
    blocks = []
    current_block = []
    footnotes = {}
    footnote_order = []  # Maintain footnote order
    print(f"DEBUG: Content length: {len(content)} characters")
    print(f"DEBUG: First 100 chars of content: {content[:100]}")

    # First pass: collect all footnote definitions
    for line in content.split('\n'):
        if line.startswith('[^'):
            match = re.match(r'\[\^(\d+)\]:\s*(.*)', line)
            if match:
                num, text = match.groups()
                footnotes[num] = text
                if num not in footnote_order:
                    footnote_order.append(num)

    # Second pass: process text while preserving footnote references
    for line in content.split('\n'):
        if line.startswith('[^'):
            continue  # Skip footnote definitions here
        elif line.startswith(('#', '>', '- ', '* ', '1.')) or line.strip().startswith('|'):
            if current_block:
                blocks.append('\n'.join(current_block))
                current_block = []
            blocks.append(line)
        elif not line.strip():
            if current_block:
                blocks.append('\n'.join(current_block))
                current_block = []
        else:
            current_block.append(line)

    if current_block:
        blocks.append('\n'.join(current_block))

    # Setup LLM - prefer academic_llm over general llm
    model_id = academic_llm if academic_llm else (llm if llm else "ollama/qwen3")
    print(f"DEBUG: Using model_id: {model_id}")

    # Ensure model_id is a string
    if not isinstance(model_id, str):
        model_id = str(model_id)
        print(f"DEBUG: Converted model_id to string: {model_id}")

    # Default to qwen3 if model_id is empty after conversion
    if not model_id.strip():
        model_id = "ollama/qwen3"
        print(f"DEBUG: Using default model_id: {model_id}")

    # For Gemini models, ensure required packages are installed
    if model_id.startswith("gemini/"):
        # We already imported google.generativeai at the top of the file
        if not os.getenv("GOOGLE_API_KEY"):
            raise ValueError("GOOGLE_API_KEY environment variable not set.")

    # Get LLM configuration and initialize
    print(f"DEBUG: Getting LLM config for model_id={model_id}")
    try:
        lm_config = get_lm_config(model_id, base_url=base_url)

        # If we got a custom LM instance directly, use it
        if "lm" in lm_config:
            lm = lm_config["lm"]
            # Set any extra parameters
            if hasattr(lm, "kwargs"):
                lm.kwargs["max_tokens"] = max_tokens
                lm.kwargs["timeout"] = timeout
                lm.kwargs["temperature"] = temperature
        else:
            # Otherwise configure a standard DSPy LM
            lm_config["max_tokens"] = max_tokens
            lm_config["timeout_s"] = timeout
            lm_config["temperature"] = temperature
            lm = dspy.LM(**lm_config)

        print(f"DEBUG: LLM config applied: {lm_config}")
    except Exception as e:
        print(f"DEBUG: Error getting LLM config: {e}")
        traceback.print_exc()
        raise

    try:
        # Configure DSPy to use this language model
        dspy.configure(lm=lm)
        print(f"Successfully configured DSPy with model: {model_id}")
    except Exception as e:
        print(f"Error initializing LLM: {e}")
        traceback.print_exc()
        raise ValueError(f"Failed to initialize LLM with model {model_id}: {str(e)}")

    # Process each block
    academic_blocks = []
    for block in blocks:
        if not block.strip():
            continue
        if block.startswith(('#', '>', '- ', '* ', '1.')) or block.strip().startswith('|'):
            academic_blocks.append(block)
            continue

        class AcademicRewrite(dspy.Signature):
            """
            Rewrite this text in formal academic style in {academic_lang}. Follow these rules strictly:
            1. Correct any basic grammar, punctuation, or usage errors.
            2. Improve clarity while preserving the original meaning and scholarly tone (trying your best not to change the structure of sentence)
            3. Maintaining the original meaning and length (97% of original)
            4. IMPORTANT: Preserve ALL footnote references (e.g., [^1], [^2]) exactly as they appear
            """
            text: str = dspy.InputField(desc=f"Text to rewrite in academic {academic_lang}")
            academic: str = dspy.OutputField(desc=f"Academic rewritten text in {academic_lang}")

        # Monkey patch the parse method in dspy to handle list responses
        original_parse = dspy.adapters.JsonAdapter.parse

        def safe_parse(self, signature, response):
            try:
                # Handle list responses by converting to a dictionary format
                if isinstance(response, list):
                    print(f"Converting list response to dictionary format")
                    if len(response) > 0 and isinstance(response[0], dict):
                        response = {"choices": [{"message": {"content": json.dumps(response[0])}}]}
                    else:
                        response = {"choices": [{"message": {"content": "No valid content found"}}]}
                return original_parse(self, signature, response)
            except Exception as e:
                print(f"Error in parse: {e}")
                # Return a basic output with all the expected fields
                return {field: "" for field in signature.output_fields}

        # Apply the monkey patch
        dspy.adapters.JsonAdapter.parse = safe_parse

        # Create a custom wrapped version of ChainOfThought that handles errors
        class SafeChainOfThought:
            def __init__(self, signature_class):
                self.chain = dspy.ChainOfThought(signature_class)
                self.max_retries = 2

            def __call__(self, **kwargs):
                # Get the input text for fallback
                input_text = kwargs.get('text', '')

                for attempt in range(self.max_retries + 1):
                    try:
                        print(f"Academic rewrite attempt {attempt+1}/{self.max_retries+1}")

                        # Call the underlying chain
                        result = self.chain(**kwargs)

                        # If we have a valid result with academic field, return it
                        if hasattr(result, 'academic') and result.academic:
                            return result.academic

                    except Exception as e:
                        print(f"Error in rewrite attempt {attempt+1}: {e}")
                        if attempt == self.max_retries:
                            print("All retry attempts failed")

                # If we got here, all attempts failed - return original text
                print("Falling back to original text")
                return input_text

        # Use our safe wrapper
        safe_academic_rewrite = SafeChainOfThought(AcademicRewrite)

        try:
            # Process each block with our safe wrapper
            rewritten_text = safe_academic_rewrite(text=block)
            academic_blocks.append(rewritten_text)
        except Exception as e:
            print(f"Error during academic rewrite: {e}")
            # Fall back to original text if rewrite fails
            academic_blocks.append(block)

    # Combine processed blocks
    academic_text = '\n\n'.join(academic_blocks)

    # Append footnotes in original order
    if footnotes:
        academic_text += '\n\n'
        for num in footnote_order:
            academic_text += f'[^{num}]: {footnotes[num]}\n'

    # Only write to file if output_dir is explicitly requested
    if output_dir:
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        out_file = os.path.join(output_dir, f"{base_name}_academic.md")
        with open(out_file, 'w', encoding='utf-8') as f:
            f.write(academic_text)
        print(f"Note: Academic text written to {out_file}")

    return academic_text


def process_docx(
    file_path,
    output_dir=None,
    llm="",
    academic_llm="",  # Added academic_llm parameter
    academic_lang="English",
    chunk_length=8,
    max_tokens=50000,
    timeout=3600,
    temperature=0.1,
    base_url="http://localhost:11434",
):
    """Process a docx file with only essential outputs:
    1. Original markdown
    2. Comparison markdown (redlines)
    """
    print(f"DEBUG: process_docx function called with file_path={file_path}, llm={llm}, academic_llm={academic_llm}")
    print(f"DEBUG: All params: output_dir={output_dir}, academic_lang={academic_lang}, chunk_length={chunk_length}, max_tokens={max_tokens}, timeout={timeout}, temperature={temperature}")

    # Type checking for parameters
    if not isinstance(file_path, str):
        print(f"WARNING: file_path is not a string: {type(file_path)}")
        file_path = str(file_path) if file_path is not None else ""

    if not isinstance(llm, str):
        print(f"WARNING: llm is not a string: {type(llm)}")
        llm = str(llm) if llm is not None else ""

    if not isinstance(academic_llm, str):
        print(f"WARNING: academic_llm is not a string: {type(academic_llm)}")
        academic_llm = str(academic_llm) if academic_llm is not None else ""

    if not isinstance(chunk_length, int):
        print(f"WARNING: chunk_length is not an integer: {type(chunk_length)}")
        try:
            chunk_length = int(chunk_length)
        except (ValueError, TypeError):
            chunk_length = 8
    # For Gemini models, check if API key is set
    if (llm and llm.startswith("gemini/")) or (academic_llm and academic_llm.startswith("gemini/")):
        if not os.getenv("GOOGLE_API_KEY"):
            raise ValueError("GOOGLE_API_KEY environment variable not set.")

    if not output_dir:
        output_dir = os.path.dirname(file_path)
    os.makedirs(output_dir, exist_ok=True)
    base_name = os.path.splitext(os.path.basename(file_path))[0]

    # 1. Convert original to markdown using pandoc to preserve footnotes
    original_md = os.path.join(output_dir, f"{base_name}_original.md")
    try:
        subprocess.run([
            'pandoc',
            '-f', 'docx',
            '-t', 'markdown',
            '--wrap=none',
            '--markdown-headings=atx',
            '--reference-links',
            file_path,
            '-o', original_md
        ], check=True)
        print(f"Created original markdown: {original_md}")
    except subprocess.CalledProcessError:
        # Fallback to simple conversion if pandoc fails
        doc = Document(file_path)
        with open(original_md, 'w', encoding='utf-8') as f:
            f.write('\n\n'.join(p.text for p in doc.paragraphs if p.text.strip()))
        print(f"Created original markdown using fallback method: {original_md}")

    # 2. Generate academic rewrite using original markdown as input
    try:
        print(f"DEBUG: About to call academic() with original_md={original_md}")
        # We already converted these to strings in the type checking at the start of the function
        academic_text = academic(
            original_md,  # Changed from file_path to original_md
            output_dir=output_dir,
            llm=llm,
            academic_llm=academic_llm,
            academic_lang=academic_lang,
            chunk_length=chunk_length,
            max_tokens=max_tokens,
            timeout=timeout,
            temperature=temperature,
            base_url=base_url,
        )
        print("DEBUG: academic() function completed successfully")
    except Exception as e:
        print(f"DEBUG: Error in academic() function: {e}")
        traceback.print_exc()
        raise

    # 3. Generate comparison markdown using redlines
    compare_md = os.path.join(output_dir, f"{base_name}_compare.md")
    with open(original_md, 'r', encoding='utf-8') as f:
        original_text = f.read()
    print(f"Creating redlines comparison between original and academic text")
    try:
        diff = Redlines(original_text, academic_text)
        print("Redlines comparison created successfully")
    except Exception as e:
        print(f"Error creating redlines comparison: {e}")
        raise
    with open(compare_md, 'w', encoding='utf-8') as f:
        f.write("# Document Comparison\n\n")
        f.write(f"**Original:** {base_name}\n")
        f.write(f"**Academic Rewrite:** {base_name}_academic\n\n")
        f.write("## Changes\n\n")
        f.write(diff.output_markdown)
    print(f"Created comparison markdown: {compare_md}")

    # Only return the essential outputs
    result = {
        'original_md': original_md,
        'compare_md': compare_md
    }
    print(f"Process complete. Returning original and comparison markdown files.")
    return result

def split_text_preserve_footnotes(text):
    """Split text into paragraphs while preserving footnote references"""
    print(f"DEBUG: split_text_preserve_footnotes called with text of length {len(text) if text else 0}")
    print(f"DEBUG: text type: {type(text)}")

    # Defensive check - ensure text is a string
    if not isinstance(text, str):
        print(f"WARNING: text is not a string: {type(text)}")
        if text is None:
            text = ""
        else:
            try:
                text = str(text)
            except Exception as e:
                print(f"ERROR converting text to string: {e}")
                text = ""

    paras = []
    current = []
    footnotes = {}

    for line in text.split('\n'):
        if line.startswith('[^'):
            # Store footnote definition
            match = re.match(r'\[\^(\d+)\]:\s*(.*)', line)
            if match:
                footnotes[match.group(1)] = match.group(2)
        elif line.strip():
            current.append(line)
        elif current:
            paras.append(('\n'.join(current), footnotes))
            current = []
            footnotes = {}

    if current:
        paras.append(('\n'.join(current), footnotes))

    return paras

def add_text_with_footnotes(paragraph, content):
    """Add text to paragraph while properly handling footnote references"""
    text, footnotes = content
    parts = re.split(r'(\[\^\d+\])', text)

    for part in parts:
        if re.match(r'\[\^(\d+)\]', part):
            # Add footnote reference
            footnote_num = re.search(r'\[\^(\d+)\]', part).group(1)
            if footnote_num in footnotes:
                paragraph.add_run().add_footnote(footnotes[footnote_num])
        else:
            paragraph.add_run(part)
